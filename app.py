# Arquivo: app.py (com cabeçalho formatado)
from flask import Flask, render_template, request, send_from_directory, url_for, redirect, session
from docx import Document
from docx.shared import Pt
import os
import datetime
import re
from spellchecker import SpellChecker

# --- CONFIGURAÇÃO DA APLICAÇÃO ---
app = Flask(__name__)
UPLOAD_FOLDER = 'orcamentos_gerados'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'agora-vai-funcionar-com-certeza-2025'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


DADOS_EMPRESA = {
    "razao_social": "NOME DA EMPRESA",
    "nome_responsavel": "NOME DO RESPONSÁVEL",
    "cnpj": "00.000.000/0001-22",
    "inscricao_municipal": "07.05.744.522-9",
    "endereco": "Endereço completo",
    "telefone": "(41) 90000-1549"
}
# ---------------------------------------------------

def gerar_orcamento_docx(dados):
    try:
        doc = Document()
        
        # --- Cabeçalho da Empresa ---
        # Adiciona cada linha de informação com a formatação correta
        p_razao = doc.add_paragraph()
        p_razao.add_run(dados["razao_social"]).bold = True

        p_nome = doc.add_paragraph()
        p_nome.add_run(dados["nome_responsavel"]).bold = True

        p_cnpj = doc.add_paragraph()
        p_cnpj.add_run('CNPJ: ').bold = True
        p_cnpj.add_run(dados["cnpj"])

        p_im = doc.add_paragraph()
        p_im.add_run('Inscrição Municipal: ').bold = True
        p_im.add_run(dados["inscricao_municipal"])

        doc.add_paragraph(dados["endereco"])
        doc.add_paragraph(dados["telefone"])
        
        doc.add_paragraph("\n" + ("-"*50) + "\n")

        titulo = doc.add_paragraph()
        run_titulo = titulo.add_run("ORÇAMENTO")
        run_titulo.bold = True
        run_titulo.font.size = Pt(18)
        titulo.alignment = 1
        doc.add_paragraph(f"\nCliente: {dados['cliente_nome']}")
        doc.add_paragraph(f"Data: {dados['data_orcamento']}\n")
        doc.add_heading("Serviços Inclusos:", level=2)
        for servico in dados["servicos_inclusos"]:
            doc.add_paragraph(servico, style='List Bullet')
        doc.add_paragraph("\n")
        paragrafo_valor = doc.add_paragraph()
        paragrafo_valor.add_run("Valor Total: ").bold = True
        paragrafo_valor.add_run(f"R$ {dados['valor_total']}")
        paragrafo_pagamento = doc.add_paragraph()
        paragrafo_pagamento.add_run("Forma de Pagamento: ").bold = True
        paragrafo_pagamento.add_run(dados['forma_pagamento'])
        nome_cliente_seguro = re.sub(r'[^a-zA-Z0-9_.-]', '_', dados['cliente_nome']).lower()
        nome_arquivo_base = f"orcamento_{nome_cliente_seguro}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
        nome_arquivo_docx = f"{nome_arquivo_base}.docx"
        caminho_completo = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo_docx)
        doc.save(caminho_completo)
        return nome_arquivo_docx
    except Exception as e:
        print(f"Ocorreu um erro ao gerar o arquivo DOCX: {e}")
        return None

@app.route('/', methods=['GET', 'POST'])
def index():
    data_hoje = datetime.date.today().strftime('%d/%m/%Y')
    return render_template('index.html', empresa=DADOS_EMPRESA, data_hoje=data_hoje)

@app.route('/verificar', methods=['POST'])
def verificar_ortografia():
    spell = SpellChecker(language='pt')
    texto_servicos = request.form['servicos_inclusos']
    palavras = re.findall(r'\b\w+\b', texto_servicos.lower())
    palavras_erradas = spell.unknown(palavras)
    
    dados_completos = {key: value for key, value in request.form.items()}

    if palavras_erradas:
        erros_com_sugestoes = {palavra: spell.correction(palavra) for palavra in palavras_erradas}
        session['dados_formulario'] = dados_completos
        session['erros_ortografia'] = erros_com_sugestoes
        return redirect(url_for('pagina_correcao'))
    else:
        return gerar_orcamento(dados_completos)

@app.route('/corrigir')
def pagina_correcao():
    dados = session.get('dados_formulario', {})
    erros = session.get('erros_ortografia', {})
    return render_template('corrigir.html', erros=erros, dados=dados)

@app.route('/preview', methods=['POST'])
def preview_correcao():
    dados_formulario = {key: value for key, value in request.form.items()}
    texto_original = dados_formulario.get('servicos_inclusos', '')
    
    spell = SpellChecker(language='pt')
    palavras = re.findall(r'\b\w+\b', texto_original.lower())
    palavras_erradas = spell.unknown(palavras)
    erros_com_sugestoes = {palavra: spell.correction(palavra) for palavra in palavras_erradas}
    
    texto_corrigido = texto_original
    if erros_com_sugestoes:
        for palavra_errada, sugestao in erros_com_sugestoes.items():
            if sugestao:
                texto_corrigido = re.sub(r'\b' + re.escape(palavra_errada) + r'\b', sugestao, texto_corrigido, flags=re.IGNORECASE)
    
    dados_formulario['servicos_inclusos'] = texto_corrigido
    return render_template('corrigir.html', erros=erros_com_sugestoes, dados=dados_formulario)

@app.route('/gerar', methods=['POST'])
def gerar_rota():
    dados_formulario = {key: value for key, value in request.form.items()}
    return gerar_orcamento(dados_formulario)

def gerar_orcamento(dados):
    dados_para_doc = dados.copy()
    dados_para_doc["servicos_inclusos"] = [s.strip().capitalize() for s in dados["servicos_inclusos"].splitlines() if s.strip()]
    
    dados_para_doc.update(DADOS_EMPRESA) 
    
    nome_docx = gerar_orcamento_docx(dados_para_doc)
    
    if nome_docx:
        return redirect(url_for('pagina_resultado', filename=nome_docx))
    else:
        return "<h1>Erro ao gerar o documento. Verifique o console do servidor.</h1>", 500

@app.route('/resultado/<filename>')
def pagina_resultado(filename):
    url_docx = url_for('download_file', filename=filename)
    return render_template('resultado.html', url_docx=url_docx)

@app.route('/downloads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
